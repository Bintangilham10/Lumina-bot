"""Tests for document loading validation."""

from __future__ import annotations

import tempfile
import unittest
import warnings
from pathlib import Path

import docx
import fitz
from ebooklib import epub

from core.loader import load_document


class LoaderTests(unittest.TestCase):
    def test_load_document_rejects_missing_file(self) -> None:
        missing_path = Path.cwd() / "missing-document.pdf"

        with self.assertRaisesRegex(FileNotFoundError, "Document not found"):
            load_document(missing_path)

    def test_load_document_rejects_unsupported_file_type(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as temp_dir:
            path = Path(temp_dir) / "notes.txt"
            path.write_text("Hello world", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "Unsupported file type"):
                load_document(path)

    def test_load_document_rejects_directory_paths(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as temp_dir:
            path = Path(temp_dir) / "book.pdf"
            path.mkdir()

            with self.assertRaisesRegex(ValueError, "not a file"):
                load_document(path)

    def test_load_document_reads_pdf_text(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as temp_dir:
            path = Path(temp_dir) / "sample.pdf"
            pdf = fitz.open()
            page = pdf.new_page()
            page.insert_text((72, 72), "Lumina PDF content.")
            pdf.save(path)
            pdf.close()

            loaded = load_document(path)

            self.assertEqual(loaded.file_type, "PDF")
            self.assertEqual(loaded.total_pages, 1)
            self.assertIn("Lumina PDF content.", loaded.documents[0].page_content)

    def test_load_document_reads_docx_paragraphs_and_tables(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as temp_dir:
            path = Path(temp_dir) / "sample.docx"
            document = docx.Document()
            document.add_paragraph("Lumina DOCX paragraph.")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Column A"
            table.rows[0].cells[1].text = "Column B"
            document.save(path)

            loaded = load_document(path)

            self.assertEqual(loaded.file_type, "DOCX")
            self.assertEqual(loaded.total_pages, 1)
            self.assertIn("Lumina DOCX paragraph.", loaded.documents[0].page_content)
            self.assertIn("Column A | Column B", loaded.documents[0].page_content)

    def test_load_document_reads_epub_content_without_navigation_page(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as temp_dir:
            path = Path(temp_dir) / "sample.epub"
            book = epub.EpubBook()
            book.set_identifier("lumina-smoke")
            book.set_title("Lumina Smoke")
            book.set_language("en")
            chapter = epub.EpubHtml(title="Chapter 1", file_name="chapter.xhtml", lang="en")
            chapter.content = (
                "<html><body><h1>Chapter</h1>"
                "<p>Lumina EPUB content.</p></body></html>"
            )
            book.add_item(chapter)
            book.toc = (chapter,)
            book.spine = ["nav", chapter]
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            epub.write_epub(str(path), book)

            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                loaded = load_document(path)

            self.assertEqual(loaded.file_type, "EPUB")
            self.assertEqual(loaded.total_pages, 1)
            self.assertIn("Lumina EPUB content.", loaded.documents[0].page_content)


if __name__ == "__main__":
    unittest.main()
