"""Document loading utilities for PDF, DOCX, and EPUB files."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import zipfile

import docx
import fitz
from bs4 import BeautifulSoup
from ebooklib import ITEM_DOCUMENT
from ebooklib import epub
from langchain_core.documents import Document

from utils.helpers import clean_text, is_supported_file, supported_extensions_text


MAX_ZIP_ENTRY_COUNT = 2000
MAX_ZIP_ENTRY_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_ZIP_TOTAL_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_ZIP_COMPRESSION_RATIO = 1000.0
MIN_RATIO_CHECK_BYTES = 1024 * 1024


@dataclass(frozen=True)
class LoadedDocument:
    """Loaded document payload and display metadata."""

    filename: str
    file_path: Path
    file_type: str
    total_pages: int
    documents: list[Document]


def load_document(file_path: str | Path) -> LoadedDocument:
    """Load a PDF, DOCX, or EPUB file into LangChain Document objects."""
    path = Path(file_path).expanduser()
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    if not path.is_file():
        raise ValueError(f"Document path is not a file: {path}")
    if not is_supported_file(path):
        raise ValueError(
            f"Unsupported file type '{path.suffix}'. Supported formats: {supported_extensions_text()}."
        )

    suffix = path.suffix.lower()
    _validate_file_signature(path, suffix)
    if suffix == ".pdf":
        documents = _load_pdf(path)
    elif suffix == ".docx":
        documents = _load_docx(path)
    else:
        documents = _load_epub(path)

    if not documents:
        raise ValueError("No readable text was found in this document.")

    return LoadedDocument(
        filename=path.name,
        file_path=path.resolve(),
        file_type=suffix.lstrip(".").upper(),
        total_pages=len(documents),
        documents=documents,
    )


def _base_metadata(path: Path, file_type: str) -> dict[str, str]:
    return {
        "source": str(path),
        "filename": path.name,
        "file_type": file_type,
    }


def _validate_file_signature(path: Path, suffix: str) -> None:
    """Reject common extension spoofing before handing files to parsers."""
    if suffix == ".pdf":
        with path.open("rb") as file:
            header = file.read(5)
        if header != b"%PDF-":
            raise ValueError("File content does not look like a PDF document.")
        return

    if suffix == ".docx":
        _validate_zip_members(path, {"[Content_Types].xml", "word/document.xml"}, "DOCX")
        return

    _validate_epub_signature(path)


def _validate_zip_members(path: Path, required_members: set[str], file_type: str) -> None:
    if not zipfile.is_zipfile(path):
        raise ValueError(f"File content does not look like a {file_type} document.")
    with zipfile.ZipFile(path) as archive:
        _validate_zip_archive_safety(archive, file_type)
        names = set(archive.namelist())
    if not required_members.issubset(names):
        raise ValueError(f"File content does not look like a {file_type} document.")


def _validate_epub_signature(path: Path) -> None:
    if not zipfile.is_zipfile(path):
        raise ValueError("File content does not look like an EPUB document.")
    with zipfile.ZipFile(path) as archive:
        _validate_zip_archive_safety(archive, "EPUB")
        names = set(archive.namelist())
        if "mimetype" not in names:
            raise ValueError("File content does not look like an EPUB document.")
        mimetype = archive.read("mimetype").decode("utf-8", errors="ignore").strip()
    if mimetype != "application/epub+zip":
        raise ValueError("File content does not look like an EPUB document.")


def _validate_zip_archive_safety(archive: zipfile.ZipFile, file_type: str) -> None:
    entries = [info for info in archive.infolist() if not info.is_dir()]
    if len(entries) > MAX_ZIP_ENTRY_COUNT:
        raise ValueError(
            f"{file_type} archive has too many files to process safely "
            f"({len(entries)} > {MAX_ZIP_ENTRY_COUNT})."
        )

    total_uncompressed = 0
    for entry in entries:
        total_uncompressed += entry.file_size
        if entry.file_size > MAX_ZIP_ENTRY_UNCOMPRESSED_BYTES:
            raise ValueError(
                f"{file_type} archive entry '{entry.filename}' is too large to process safely."
            )
        compressed_size = max(entry.compress_size, 1)
        compression_ratio = entry.file_size / compressed_size
        if (
            entry.file_size >= MIN_RATIO_CHECK_BYTES
            and compression_ratio > MAX_ZIP_COMPRESSION_RATIO
        ):
            raise ValueError(
                f"{file_type} archive entry '{entry.filename}' has an unsafe compression ratio."
            )

    if total_uncompressed > MAX_ZIP_TOTAL_UNCOMPRESSED_BYTES:
        raise ValueError(
            f"{file_type} archive expands beyond the safe processing limit "
            f"({total_uncompressed} bytes > {MAX_ZIP_TOTAL_UNCOMPRESSED_BYTES} bytes)."
        )


def _load_pdf(path: Path) -> list[Document]:
    documents: list[Document] = []
    metadata = _base_metadata(path, "PDF")

    with fitz.open(path) as pdf:
        for index, page in enumerate(pdf, start=1):
            text = clean_text(page.get_text("text"))
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={**metadata, "page": index, "section": f"Page {index}"},
                    )
                )

    return documents


def _load_docx(path: Path) -> list[Document]:
    document = docx.Document(path)
    paragraphs = [clean_text(paragraph.text) for paragraph in document.paragraphs]
    table_rows: list[str] = []

    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                table_rows.append(row_text)

    text_parts = [paragraph for paragraph in paragraphs if paragraph] + table_rows
    text = "\n\n".join(text_parts)

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={**_base_metadata(path, "DOCX"), "page": 1, "section": "Document"},
        )
    ]


def _load_epub(path: Path) -> list[Document]:
    book = epub.read_epub(str(path))
    documents: list[Document] = []
    metadata = _base_metadata(path, "EPUB")

    for item in book.get_items_of_type(ITEM_DOCUMENT):
        if _is_epub_navigation_item(item):
            continue

        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = clean_text(soup.get_text(separator="\n"))
        if text:
            index = len(documents) + 1
            title = item.get_name() or f"Section {index}"
            documents.append(
                Document(
                    page_content=text,
                    metadata={**metadata, "page": index, "section": title},
                )
            )

    return documents


def _is_epub_navigation_item(item) -> bool:
    """Return whether an EPUB document item is navigation-only content."""
    if isinstance(item, epub.EpubNav):
        return True
    name = str(item.get_name() or "").lower()
    return name in {"nav.xhtml", "toc.xhtml", "toc.html"}
